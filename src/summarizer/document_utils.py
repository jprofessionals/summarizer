from PyPDF2 import PdfReader
import docx
from typing import Dict, List


def extract_text_from_txt(txt_file) -> str:
    """
    Extracts text from a TXT file.

    Args:
        txt_file (file-like object): The text file to extract text from.

    Returns:
        str: The extracted text.
    """
    try:
        text = txt_file.read().decode("utf-8")
    except Exception as e:
        raise ValueError(f"Error reading TXT file: {e}")
    return text


def extract_text_from_pdf(pdf_file) -> str:
    """
    Extracts text from a PDF file.

    Args:
        pdf_file (file-like object): The PDF file to extract text from.

    Returns:
        str: The extracted text.
    """
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {e}")
    return text


def extract_text_from_docx(docx_file) -> str:
    """
    Extracts text from a DOCX file.

    Args:
        docx_file (file-like object): The DOCX file to extract text from.

    Returns:
        str: The extracted text.
    """
    try:
        doc = docx.Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {e}")
    return text


def extract_requirements_from_docx(docx_file) -> Dict[str, List[str]]:
    """
    Extracts matrix-like data from the requirements DOCX file.

    Args:
        docx_file (file-like object): The DOCX file to extract requirements from.

    Returns:
        Dict[str, List[str]]: A dictionary of requirements extracted from the DOCX file.
    """
    requirements: Dict[str, List[str]] = {}
    doc = docx.Document(docx_file)
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = row.cells
            for i, header in enumerate(headers):
                key = header
                value = cells[i].text.strip()
                if key and value:
                    if key not in requirements:
                        requirements[key] = []
                    requirements[key].append(value)
    return requirements


def format_requirements_for_display(requirements_dict: Dict[str, List[str]]) -> str:
    """
    Formats the requirements dictionary for display, adding numbering for each line.

    Args:
        requirements_dict (Dict[str, List[str]]): The dictionary of requirements.

    Returns:
        str: The formatted string of requirements.
    """
    formatted_requirements = []
    counter = 1
    for key, values in requirements_dict.items():
        for value in values:
            formatted_requirements.append(f"{counter}. {key}: {value}")
            counter += 1
    return "\n".join(formatted_requirements)
